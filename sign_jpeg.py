import os
from docx import Document
from docx.shared import Inches
import fitz  # PyMuPDF
from PIL import Image, ImageDraw

def add_image_to_docx(docx_path, image_path):
    # Load the DOCX document
    doc = Document(docx_path)

    # Access the last paragraph on the last page
    last_paragraph = doc.paragraphs[-1]

    # Add the image to the last paragraph
    last_paragraph.add_run().add_picture(image_path, width=Inches(2))

    # Save the modified document
    new_docx_path = docx_path.replace('.docx', '_modified.docx')
    doc.save(new_docx_path)
    return new_docx_path

def add_image_to_pdf(pdf_path, image_path):
    # Load the PDF
    pdf = fitz.open(pdf_path)

    # Load the image and convert to a Pixmap
    image = Image.open(image_path)
    image_bytes = io.BytesIO()
    image.save(image_bytes, format='PNG')
    image_bytes.seek(0)
    pix = fitz.Pixmap(image_bytes.getvalue())

    # Get the last page
    last_page = pdf[-1]
    rect = last_page.rect
    # Define the position for the image (bottom right corner)
    # Assuming we want to place the image at the bottom right with a size 100x100 points
    x = rect.width - 100
    y = rect.height - 100
    last_page.insert_image(fitz.Rect(x, y, x + 100, y + 100), pixmap=pix)

    # Save the modified PDF
    new_pdf_path = pdf_path.replace('.pdf', '_modified.pdf')
    pdf.save(new_pdf_path)
    pdf.close()
    return new_pdf_path

def add_image_to_jpeg(jpeg_path, image_path):
    # Load the JPEG image
    original = Image.open(jpeg_path)
    original = original.convert("RGBA")

    # Load the image to be added
    image = Image.open(image_path)
    image = image.convert("RGBA")

    # Calculate the position to place the image (bottom right corner)
    x = original.width - image.width
    y = original.height - image.height

    # Create a new image by combining the original and the image to be added
    combined = Image.new("RGBA", original.size)
    combined.paste(original, (0, 0))
    combined.paste(image, (x, y), image)

    # Save the combined image
    new_jpeg_path = jpeg_path.replace('.jpg', '_modified.jpg')
    combined = combined.convert("RGB")  # Convert back to RGB before saving
    combined.save(new_jpeg_path, 'JPEG')
    return new_jpeg_path

def add_image_to_file(file_path, image_path):
    # Get the file extension
    _, ext = os.path.splitext(file_path)

    # Check the file extension and call the appropriate function
    if ext.lower() == '.docx':
        return add_image_to_docx(file_path, image_path)
    elif ext.lower() == '.pdf':
        return add_image_to_pdf(file_path, image_path)
    elif ext.lower() in ['.jpeg', '.jpg']:
        return add_image_to_jpeg(file_path, image_path)
    else:
        raise ValueError("Unsupported file format. Please provide a DOCX, PDF, or JPEG file.")

# Example usage
file_path = 'example.jpg' #'example.docx'  # or 'example.pdf' or 'example.jpeg'
image_path = 'image.png'

try:
    modified_file_path = add_image_to_file(file_path, image_path)
    print(f"Modified file saved as: {modified_file_path}")
except ValueError as e:
    print(e)